from docx import Document


def extract_text_from_docx(filepath):
    """
    Extract all text from a .docx file, including both paragraphs and tables.

    ROOT CAUSE FIX: The original function only read doc.paragraphs, which
    returned empty text because this document stores all invoice data inside
    table cells. Tables must be read separately.
    """
    doc = Document(filepath)
    parts = []

    # Read top-level paragraphs (usually empty in this doc, but kept for safety)
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Read all table cells — this is where the invoice data actually lives
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    return parts  # Returns a list of invoice text blocks (one per cell)


def split_invoices(text):
    """
    ROOT CAUSE FIX: Each table cell is already one invoice — no splitting needed.
    'text' is now a list of strings returned by extract_text_from_docx.
    We just filter out any cells that are too short to be a real invoice.
    """
    invoices = []
    for chunk in text:
        if len(chunk.strip()) > 100:
            invoices.append(chunk.strip())
    return invoices